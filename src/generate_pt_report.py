"""
PT Quarterly Report Generator — A&G Physiotherapy Inc.
Usage:  python generate_pt_report.py YourFile.xlsx
Output: PT_Q1_2026_Report.docx  (same folder as script)
Needs:  pip install pandas openpyxl python-docx matplotlib pillow
"""
import sys, os, io, subprocess, re
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x1B,0x3A,0x6B); TEAL  = RGBColor(0x1A,0x7A,0x8A)
WHITE = RGBColor(0xFF,0xFF,0xFF); DARK  = RGBColor(0x2C,0x3E,0x50)
MID   = RGBColor(0x7F,0x8C,0x8D); GREEN = RGBColor(0x1E,0x84,0x49)
AMBER = RGBColor(0xB7,0x77,0x0D); RED   = RGBColor(0x92,0x2B,0x21)

H = dict(navy="1B3A6B",teal="1A7A8A",white="FFFFFF",dark="2C3E50",mid="7F8C8D",
         lblue="EBF5FB",lgreen="EAFAF1",lyellow="FEFAED",lred="FDEDEC",
         green="1E8449",amber="B7770D",red="922B21",bg="F4F8FC")
MC = dict(navy="#1B3A6B",teal="#1A7A8A",green="#27AE60",amber="#E67E22",
          red="#C0392B",blue="#2980B9",lblue="#AED6F1",bg="#F4F8FC")
TW = 9360

def _new(tag,**attrs):
    el=OxmlElement(tag)
    for k,v in attrs.items(): el.set(qn(k),str(v))
    return el

def _tcPr(cell): return cell._tc.get_or_add_tcPr()

def cell_setup(cell,w,fill=None,margins=(80,80,120,120),borders=None,valign=None):
    tcPr=_tcPr(cell)
    tcPr.insert(0,_new("w:tcW",**{"w:w":str(w),"w:type":"dxa"}))
    if fill: tcPr.append(_new("w:shd",**{"w:val":"clear","w:color":"auto","w:fill":fill}))
    if borders:
        b=OxmlElement("w:tcBorders")
        for side,val in borders.items():
            el=OxmlElement(f"w:{side}")
            if val is None: el.set(qn("w:val"),"nil")
            else:
                el.set(qn("w:val"),val.get("style","single"))
                el.set(qn("w:sz"),str(val.get("sz",4)))
                el.set(qn("w:color"),val.get("color","DDDDDD"))
            b.append(el)
        tcPr.append(b)
    if margins:
        m=OxmlElement("w:tcMar")
        for side,v in zip(["top","bottom","left","right"],margins):
            el=OxmlElement(f"w:{side}"); el.set(qn("w:w"),str(v)); el.set(qn("w:type"),"dxa"); m.append(el)
        tcPr.append(m)
    if valign: tcPr.append(_new("w:vAlign",**{"w:val":valign}))

NB   = lambda: dict(top=None,bottom=None,left=None,right=None)
THIN = lambda c="DDDDDD": {"style":"single","sz":4,"color":c}
ALLB = lambda c="DDDDDD": dict(top=THIN(c),bottom=THIN(c),left=THIN(c),right=THIN(c))

def set_tbl_w(table,w):
    tblPr=table._tbl.find(qn("w:tblPr"))
    if tblPr is None: tblPr=OxmlElement("w:tblPr"); table._tbl.insert(0,tblPr)
    tblPr.append(_new("w:tblW",**{"w:w":str(w),"w:type":"dxa"}))
    lay=OxmlElement("w:tblLayout"); lay.set(qn("w:type"),"fixed"); tblPr.append(lay)

def sp(p,before=0,after=40): p._p.get_or_add_pPr().append(_new("w:spacing",**{"w:before":str(before),"w:after":str(after)}))
def nos(p): sp(p,0,0)

def rn(p,text,*,bold=False,sz=9,color=None,italic=False):
    r=p.add_run(text); r.font.name="Arial"; r.font.size=Pt(sz)
    r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color or DARK; return r

def mtbl(doc,cols,widths):
    t=doc.add_table(rows=1,cols=cols); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    set_tbl_w(t,sum(widths)); return t

def gap(doc,pts=50):
    p=doc.add_paragraph(); sp(p,0,pts); return p

def sec(doc,label):
    p=doc.add_paragraph(); sp(p,90,60)
    r=p.add_run(label.upper())
    r.font.name="Arial"; r.font.size=Pt(10); r.font.bold=True; r.font.color.rgb=WHITE
    pPr=p._p.get_or_add_pPr()
    pPr.append(_new("w:shd",**{"w:val":"clear","w:fill":H["navy"]}))
    pPr.append(_new("w:ind",**{"w:left":"120"}))

def btext(doc,text,sz=8.5,color=None,before=0,after=36):
    p=doc.add_paragraph(); sp(p,before,after)
    for i,seg in enumerate(text.split("**")):
        if seg: rn(p,seg,bold=(i%2==1),sz=sz,color=color or DARK)

def figbuf(fig):
    buf=io.BytesIO()
    fig.savefig(buf,format="png",dpi=160,bbox_inches="tight",facecolor=fig.get_facecolor())
    plt.close(fig); buf.seek(0); return buf

plt.rcParams.update({"font.family":"DejaVu Sans","axes.spines.top":False,"axes.spines.right":False})

# ── Charts ────────────────────────────────────────────────────────────────────
def ch_flow(d):
    s,a,e=d["start"],d["admissions"],d["end"]
    dec=d["deceased"]
    other=d["moved_out"]+d["non_compliant"]+d["goal_achieved"]
    cats=["Opening\nCensus","Admissions","Deaths /\nPalliative","Other\nDischarged","Closing\nCensus"]
    hts=[s,a,dec,other,e]; bots=[0,s,s+a-dec,s+a-dec-other,0]
    cols=[MC["navy"],MC["green"],MC["red"],MC["amber"],MC["teal"]]
    lbls=[str(s),f"+{a}",f"−{dec}",f"−{other}",str(e)]
    fig,ax=plt.subplots(figsize=(5.6,2.9)); fig.patch.set_facecolor(MC["bg"]); ax.set_facecolor(MC["bg"])
    for i,(h,b,c,l) in enumerate(zip(hts,bots,cols,lbls)):
        ax.bar(i,h,bottom=b,color=c,width=0.52,linewidth=0,zorder=3)
        ax.text(i,b+h+0.5,l,ha="center",va="bottom",fontsize=9,fontweight="bold",color=c)
    ax.set_xticks(range(5)); ax.set_xticklabels(cats,fontsize=8)
    ax.set_ylabel("Residents",fontsize=8.5,color="#555"); ax.set_ylim(0,(s+a)*1.22)
    ax.yaxis.grid(True,color="#ddd",zorder=0); ax.set_axisbelow(True)
    ax.spines["left"].set_color("#ccc"); ax.spines["bottom"].set_color("#ccc")
    ax.set_title("Quarterly Resident Flow",fontsize=10,fontweight="bold",color=MC["navy"],pad=8)
    fig.tight_layout(); return figbuf(fig)

def ch_census(monthly):
    months=list(monthly.keys()); vals=list(monthly.values())
    fig,ax=plt.subplots(figsize=(3.5,2.9)); fig.patch.set_facecolor(MC["bg"]); ax.set_facecolor(MC["bg"])
    bars=ax.bar(months,vals,color=MC["navy"],width=0.45,linewidth=0,zorder=3)
    for bar,v in zip(bars,vals):
        ax.text(bar.get_x()+bar.get_width()/2,v+0.2,str(v),
                ha="center",va="bottom",fontsize=10,fontweight="bold",color=MC["navy"])
    ax.set_ylim(0,max(vals)*1.25); ax.yaxis.grid(True,color="#ddd",zorder=0); ax.set_axisbelow(True)
    ax.spines["left"].set_color("#ccc"); ax.spines["bottom"].set_color("#ccc")
    ax.set_ylabel("Residents",fontsize=8.5,color="#555")
    ax.set_title("Monthly Census",fontsize=10,fontweight="bold",color=MC["navy"],pad=8)
    fig.tight_layout(); return figbuf(fig)

def ch_refass(ref_m,ass_m):
    months=["Jan","Feb","Mar"]
    rv=[ref_m[m] for m in months]; av=[ass_m[m] for m in months]
    x=np.arange(3); w=0.35
    fig,ax=plt.subplots(figsize=(5.6,2.8)); fig.patch.set_facecolor(MC["bg"]); ax.set_facecolor(MC["bg"])
    b1=ax.bar(x-w/2,rv,w,color=MC["navy"],label="Referrals",linewidth=0,zorder=3)
    b2=ax.bar(x+w/2,av,w,color=MC["teal"],label="Assessments",linewidth=0,zorder=3)
    for bar,v in zip(b1,rv):
        ax.text(bar.get_x()+bar.get_width()/2,v+0.3,str(v),ha="center",va="bottom",fontsize=9,fontweight="bold",color=MC["navy"])
    for bar,v in zip(b2,av):
        ax.text(bar.get_x()+bar.get_width()/2,v+0.3,str(v),ha="center",va="bottom",fontsize=9,fontweight="bold",color=MC["teal"])
    ax.set_xticks(x); ax.set_xticklabels(months,fontsize=9.5)
    ax.set_ylabel("Count",fontsize=8.5,color="#555"); ax.set_ylim(0,max(max(rv),max(av))*1.35)
    ax.yaxis.grid(True,color="#ddd",zorder=0); ax.set_axisbelow(True)
    ax.spines["left"].set_color("#ccc"); ax.spines["bottom"].set_color("#ccc")
    ax.legend(fontsize=8,frameon=False,loc="upper left")
    ax.set_title("Monthly Referrals and Assessments",fontsize=10,fontweight="bold",color=MC["navy"],pad=8)
    fig.tight_layout(); return figbuf(fig)

def ch_programs(d):
    labels=["Ambulation +\nStrength/Balance","Chest Physio /\nPain Modality","Strengthening\n+ ROM","AAROM / PROM"]
    vals=[d["ambulation"],d["wt_bearing"],d["strengthening"],d["arom_prom"]]
    colors=[MC["navy"],MC["teal"],MC["blue"],MC["lblue"]]
    fig,ax=plt.subplots(figsize=(3.8,2.8)); fig.patch.set_facecolor(MC["bg"]); ax.set_facecolor(MC["bg"])
    _,_,ats=ax.pie(vals,colors=colors,autopct="%1.0f%%",startangle=90,
                   wedgeprops=dict(width=0.55,edgecolor="white",linewidth=1.5),pctdistance=0.76)
    for at in ats: at.set_fontsize(8.5); at.set_fontweight("bold"); at.set_color("white")
    ax.legend(labels,loc="lower center",bbox_to_anchor=(0.5,-0.24),ncol=2,fontsize=7.5,frameon=False)
    ax.set_title(f"PT Programs ({sum(vals)} residents)",fontsize=9.5,fontweight="bold",color=MC["navy"],pad=6)
    fig.tight_layout(); return figbuf(fig)

def ch_minutes(d):
    cats=["1:1 Therapy","Evaluations"]; vals=[d["one2one"],d["evaluation"]]; cols=[MC["navy"],MC["teal"]]
    fig,ax=plt.subplots(figsize=(3.8,2.2)); fig.patch.set_facecolor(MC["bg"]); ax.set_facecolor(MC["bg"])
    bars=ax.barh(cats,vals,color=cols,height=0.38,linewidth=0,zorder=3)
    for bar,v in zip(bars,vals):
        ax.text(v+150,bar.get_y()+bar.get_height()/2,f"{v:,} min",va="center",fontsize=9,fontweight="bold",color="#333")
    ax.set_title("Therapy Minutes",fontsize=9.5,fontweight="bold",color=MC["navy"],pad=8)
    ax.set_xlim(0,max(vals)*1.3); ax.xaxis.set_ticklabels([])
    ax.spines["left"].set_color("#ccc"); ax.spines["bottom"].set_visible(False)
    ax.xaxis.grid(True,color="#eee",zorder=0); ax.set_axisbelow(True); ax.tick_params(axis="y",labelsize=9)
    fig.tight_layout(); return figbuf(fig)

def ch_workforce(d):
    roles=["PTA","PT"]; g=[d["pta_group"],0]; o=[d["pta_1on1"],0]; p=[0,d["pt_hours"]]
    totals=[d["pta_total"],d["pt_hours"]]
    fig,ax=plt.subplots(figsize=(3.8,2.2)); fig.patch.set_facecolor(MC["bg"]); ax.set_facecolor(MC["bg"])
    x=np.arange(2)
    ax.bar(x,g,color=MC["teal"],width=0.45,label=f"Group ({d['pta_group']}h)",linewidth=0,zorder=3)
    ax.bar(x,o,color=MC["navy"],width=0.45,bottom=g,label=f"1:1 ({d['pta_1on1']}h)",linewidth=0,zorder=3)
    ax.bar(x,p,color=MC["blue"],width=0.45,bottom=[a+b for a,b in zip(g,o)],
           label=f"PT ({d['pt_hours']}h)",linewidth=0,zorder=3)
    for xi,tot in zip(x,totals):
        ax.text(xi,tot+0.4,f"{tot}h/wk",ha="center",fontsize=9,fontweight="bold",color="#333")
    ax.set_xticks(x); ax.set_xticklabels(roles,fontsize=10)
    ax.set_ylabel("Hours / Week",fontsize=8.5,color="#555"); ax.set_ylim(0,max(totals)*1.30)
    ax.yaxis.grid(True,color="#ddd",zorder=0); ax.set_axisbelow(True)
    ax.spines["left"].set_color("#ccc"); ax.spines["bottom"].set_color("#ccc")
    ax.legend(fontsize=7.5,frameon=False,loc="upper right")
    ax.set_title("Workforce Hours / Week",fontsize=9.5,fontweight="bold",color=MC["navy"],pad=8)
    fig.tight_layout(); return figbuf(fig)

# ── Excel reader ──────────────────────────────────────────────────────────────
def read_xl(path):
    xl=pd.read_excel(path,sheet_name=None,header=0)
    rf=xl["Resident Flow"].iloc[1]
    flow=dict(start=int(rf["Start Residents"]),admissions=int(rf["Admissions"]),
              deceased=int(rf["Discharged(Total)"]),moved_out=int(rf["Unnamed: 4"]),
              non_compliant=int(rf["Unnamed: 5"]),goal_achieved=int(rf["Unnamed: 6"]),
              end=int(rf["final(Current residents)"]))
    flow["discharged_total"]=(flow["deceased"]+flow["moved_out"]
                              +flow["non_compliant"]+flow["goal_achieved"])
    th=xl["Therapy Minutes"].iloc[0]
    therapy=dict(one2one=int(th["1:1 Minutes"]),evaluation=int(th["Evaluation Minutes"]),
                 group_sessions=int(th["Group Sessions (per week)"]))
    pr=xl["PT Programs"].iloc[0]
    programs=dict(ambulation=int(pr["Ambulation + Strength/Balance"]),wt_bearing=int(pr["Chest Physio / Pain Modality"]),
                  arom_prom=int(pr["AAROM/PROM"]),strengthening=int(pr["Strengthening + ROM"]))
    ra=xl["Referals"].iloc[0]
    referrals=dict(total=int(ra["Total Referrals"]),
                   monthly={"Jan":int(ra["Jan"]),"Feb":int(ra["Feb"]),"Mar":int(ra["Mar"])})
    aa=xl["Assesments"].iloc[0]
    assessments=dict(total=int(aa["Total Assessments"]),
                     monthly={"Jan":int(aa["Jan"]),"Feb":int(aa["Feb"]),"Mar":int(aa["Mar"])})
    st=xl["Staffing"].iloc[0]
    staffing=dict(pta_total=float(st["PTA Hours (Total)"]),pta_1on1=float(st["PTA 1:1 Hours"]),
                  pta_group=float(st["PTA Group Hours"]),pt_hours=float(st["PT Hours"]))
    sm=xl["Summary Metrics"].iloc[0]
    summary=dict(pct_1on1=float(sm["% Residents on 1:1 PT"]),total_res=int(sm["Total Residents"]))
    tb=xl["total beds"]
    total_beds=int(tb.columns[1])   # column header holds the facility capacity
    monthly_census={}
    for _,row in tb.iterrows():
        m=str(row.iloc[0]).strip(); v=row.iloc[1]
        if pd.notna(v) and m not in ["nan",""]: monthly_census[m]=int(v)
    quarter=str(xl["Resident Flow"]["Quarter"].iloc[1])
    return flow,therapy,programs,referrals,assessments,staffing,summary,monthly_census,quarter,total_beds

# ── PDF export ────────────────────────────────────────────────────────────────
def to_pdf(docx_path):
    pdf_path = docx_path.replace(".docx", ".pdf")
    out_dir  = os.path.dirname(docx_path)
    # 1) docx2pdf (uses MS Word on Mac/Win)
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return pdf_path
    except ImportError:
        pass
    # 2) LibreOffice headless
    for lo in ["/Applications/LibreOffice.app/Contents/MacOS/soffice",
               "/usr/local/bin/libreoffice", "libreoffice", "soffice"]:
        try:
            r = subprocess.run(
                [lo, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
                capture_output=True, timeout=60)
            if r.returncode == 0:
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    print("  PDF skipped — install 'docx2pdf' (pip) or LibreOffice to enable PDF export.")
    return None

# ── Main ──────────────────────────────────────────────────────────────────────
def generate(excel_path, home_name="Burton Manor, Brampton"):
    flow,therapy,programs,referrals,assessments,staffing,summary,monthly_census,quarter,total_beds=read_xl(excel_path)
    mortality_pct=round(flow["deceased"]/flow["start"]*100,1)
    hrs_1on1=round(therapy["one2one"]/60,1); hrs_eval=round(therapy["evaluation"]/60,1)
    active_res=round(summary["pct_1on1"]/100*summary["total_res"])
    prog_total=sum(programs.values())

    print("Generating charts...")
    b_flow=ch_flow(flow); b_census=ch_census(monthly_census)
    b_ra=ch_refass(referrals["monthly"],assessments["monthly"])
    b_prog=ch_programs(programs); b_min=ch_minutes(therapy); b_work=ch_workforce(staffing)

    print("Building document...")
    doc=Document()
    for s in doc.sections:
        s.top_margin=Inches(0.45); s.bottom_margin=Inches(0.45)
        s.left_margin=Inches(0.55); s.right_margin=Inches(0.55)
    doc.styles["Normal"].font.name="Arial"; doc.styles["Normal"].font.size=Pt(9)

    LOGO=os.path.join(os.path.dirname(os.path.abspath(__file__)),"logo_clean.png")

    # ── REPORT HEADER ─────────────────────────────────────────────────────────
    # 3-col: white logo cell | teal accent strip | navy title cell
    # Header: white logo cell (left) | navy title cell (right)
    # Teal right-border on logo cell acts as the visual divider — no third column
    HDR_L=2600; HDR_R=6760   # total = 9360
    t=mtbl(doc,2,[HDR_L,HDR_R])
    lc=t.cell(0,0); rc=t.cell(0,1)
    logo_bdr=dict(top=None,bottom=None,left=None,
                  right={"style":"single","sz":24,"color":H["teal"]})
    cell_setup(lc,HDR_L,fill="FFFFFF",margins=(90,90,180,180),borders=logo_bdr,valign="center")
    cell_setup(rc,HDR_R,fill=H["navy"],margins=(120,120,260,200),borders=NB(),valign="center")
    # Logo — sized to fit neatly in logo cell
    if os.path.exists(LOGO):
        p=lc.paragraphs[0]; nos(p); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(LOGO,width=Inches(1.55),height=Inches(0.37))
    else:
        p=lc.paragraphs[0]; nos(p)
        rn(p,"A & G Physiotherapy Inc.",bold=True,sz=10,color=RGBColor(0x00,0xCC,0xCC))
    # Title — 12pt keeps it on ONE line within 6760 DXA cell
    p1=rc.paragraphs[0]; sp(p1,10,4)
    rn(p1,"QUARTERLY PHYSIOTHERAPY REPORT",bold=True,sz=12,color=WHITE)
    # Teal accent rule
    p_rule=rc.add_paragraph(); sp(p_rule,6,6)
    rule_run=p_rule.add_run()
    rule_run.font.name="Arial"; rule_run.font.size=Pt(1)
    pPr=p_rule._p.get_or_add_pPr()
    pBdr=OxmlElement("w:pBdr")
    bot=OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
    bot.set(qn("w:space"),"0"); bot.set(qn("w:color"),H["teal"])
    pBdr.append(bot); pPr.append(pBdr)
    # Subtitle
    p2=rc.add_paragraph(); sp(p2,8,10)
    rn(p2,f"{home_name}   |   {quarter}",sz=8.5,color=RGBColor(0xA9,0xCC,0xE3))
    gap(doc,46)

    # ── KPIs ─────────────────────────────────────────────────────────────────
    sec(doc,"Key Performance Indicators"); gap(doc,44)

    def krow(tiles):
        n=len(tiles); cw=TW//n; t=mtbl(doc,n,[cw]*n)
        for i,(val,lbl,note,fill,badge,bcol) in enumerate(tiles):
            c=t.cell(0,i); cell_setup(c,cw,fill=fill,margins=(90,90,100,100),borders=ALLB(),valign="center")
            p1=c.paragraphs[0]; nos(p1); p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            rn(p1,str(val),bold=True,sz=20,color=NAVY)
            p2=c.add_paragraph(); nos(p2); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            rn(p2,lbl,bold=True,sz=8,color=DARK)
            p3=c.add_paragraph(); sp(p3,6,0); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
            rn(p3,note,sz=7.5,color=MID,italic=True)
            if badge:
                p4=c.add_paragraph(); sp(p4,8,0); p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
                rn(p4,badge,bold=True,sz=7.5,color=RGBColor.from_string(bcol))

    # Row 1 — 5 tiles: census, 1:1 rx, discharged total (with breakdown), non-compliant, referrals
    krow([
        (str(flow["end"]),  "Current Census",
         f"{flow['end']} of {total_beds} beds occupied", H["lblue"],
         f"↓ {flow['start']-flow['end']} from opening", H["mid"]),
        (f"{summary['pct_1on1']}%", "On Active 1:1 Rx",
         f"Of {summary['total_res']} current residents", H["lgreen"],
         "Strong coverage", H["green"]),
        (str(flow["discharged_total"]), "Discharged Total",
         f"Palliative {flow['deceased']}  ·  Moved {flow['moved_out']}  ·  Non-Comp. {flow['non_compliant']}  ·  Goal {flow['goal_achieved']}",
         H["lred"], f"{mortality_pct}% of opening census", H["red"]),
        (str(flow["non_compliant"]), "Non-Compliant / Refusal",
         "Declined or refused PT services", H["lyellow"],
         None, None),
        (str(referrals["total"]), "Total Referrals",
         f"{quarter}", H["lblue"],
         "Completed this quarter", H["amber"]),
    ])
    gap(doc,14)
    # Row 2 — 5 tiles: assessments, 1:1 min, eval min, PTA hrs, PT hrs
    krow([
        (str(assessments["total"]), "Total Assessments",
         f"{quarter}", H["lgreen"], "Completed this quarter", H["green"]),
        (f"{therapy['one2one']:,}", "1:1 Therapy Min.",
         "Direct therapy delivered", H["lblue"], None, None),
        (f"{therapy['evaluation']:,}", "Evaluation Min.",
         "Assessment time", H["lblue"], None, None),
        (f"{staffing['pta_total']}h", "PTA Hours / Week",
         "2 PTAs (1 FT + 1 PT)", H["lblue"], None, None),
        (f"{staffing['pt_hours']}h", "PT Hours / Week",
         "Mon, Tue, Thu", H["lyellow"], "3 days / week", H["amber"]),
    ])
    gap(doc,50)

    # ── RESIDENT FLOW ─────────────────────────────────────────────────────────
    sec(doc,f"Resident Flow — {quarter}"); gap(doc,44)
    t2=mtbl(doc,2,[5400,3960])
    lc2=t2.cell(0,0); rc2=t2.cell(0,1)
    cell_setup(lc2,5400,borders=NB(),margins=(30,30,0,50)); cell_setup(rc2,3960,borders=NB(),margins=(30,30,50,0))
    p=lc2.paragraphs[0]; nos(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    b_flow.seek(0); p.add_run().add_picture(b_flow,width=Inches(3.6),height=Inches(2.0))
    p=rc2.paragraphs[0]; nos(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    b_census.seek(0); p.add_run().add_picture(b_census,width=Inches(2.6),height=Inches(2.0))
    gap(doc,16)
    btext(doc,
        f"The quarter opened with **{flow['start']} residents** and closed at **{flow['end']}**. "
        f"**{flow['admissions']} new admissions** were received. "
        f"**{flow['deceased']} residents** passed away or moved to palliative care ({mortality_pct}% of opening census), "
        f"and **{flow['moved_out']}** moved out. "
        f"Monthly census: **Jan {list(monthly_census.values())[0]}**, "
        f"**Feb {list(monthly_census.values())[1]}**, **Mar {list(monthly_census.values())[2]}** residents.",
        sz=8.5,color=MID)
    gap(doc,28)

    # ── SERVICE HIGHLIGHTS — single full-width column ─────────────────────────
    shl_bullets = [
        ("● ", "A & G Physiotherapy Inc. provides ADP (Assistive Devices Program) services to eligible residents."),
        ("● ", "PT assesses all new residents on admission and completes Quarterly Assessments. The team participates in Daily Huddles, Falls Meetings, MDS, PT + NRCC Meetings and Care Conferences, and follows up on all referrals as required."),
    ]
    shl_t = mtbl(doc, 1, [TW])
    shl_c = shl_t.cell(0, 0)
    cell_setup(shl_c, TW, fill=H["lblue"], margins=(120, 120, 200, 200),
               borders=dict(top=THIN(H["teal"]), bottom=THIN(H["teal"]),
                            left=THIN(H["teal"]), right=THIN(H["teal"])))
    for idx, (dot, txt) in enumerate(shl_bullets):
        p = shl_c.paragraphs[0] if idx == 0 else shl_c.add_paragraph()
        sp(p, 0 if idx == 0 else 10, 0)
        rn(p, dot, bold=True, sz=9, color=TEAL)
        rn(p, txt, sz=8.5, color=DARK)
    gap(doc, 50)

    # ── PT REFERRALS AND ASSESSMENT — page 2 ─────────────────────────────────
    # Page break
    pb=doc.add_paragraph(); nos(pb)
    pb.add_run().add_break(__import__("docx.enum.text",fromlist=["WD_BREAK"]).WD_BREAK.PAGE)

    sec(doc,"PT Referrals and Assessment"); gap(doc,44)

    def hc(cell,text,w,fill=H["navy"],tcol=WHITE):
        cell_setup(cell,w,fill=fill,margins=(80,80,90,90),borders=ALLB("BBBBBB"),valign="center")
        p=cell.paragraphs[0]; nos(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rn(p,text,bold=True,sz=8.5,color=RGBColor.from_string(tcol) if isinstance(tcol,str) else tcol)
    def dc(cell,text,w,fill="FFFFFF",bold=False,tcol=None):
        cell_setup(cell,w,fill=fill,margins=(80,80,90,90),borders=ALLB("DDDDDD"),valign="center")
        p=cell.paragraphs[0]; nos(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rn(p,str(text),bold=bold,sz=9,color=tcol or DARK)

    ref_row=[referrals["monthly"]["Jan"],referrals["monthly"]["Feb"],referrals["monthly"]["Mar"],referrals["total"]]
    ass_row=[assessments["monthly"]["Jan"],assessments["monthly"]["Feb"],assessments["monthly"]["Mar"],assessments["total"]]

    # Outer 2-col: chart left (5040) | breakdown table right (4320) = 9360
    CW=5040; TBW=4320
    outer=mtbl(doc,2,[CW,TBW])
    chart_cell=outer.cell(0,0); tbl_cell=outer.cell(0,1)
    cell_setup(chart_cell,CW,borders=NB(),margins=(0,20,0,60),valign="center")
    cell_setup(tbl_cell,TBW,borders=NB(),margins=(20,20,80,0),valign="center")

    pc=chart_cell.paragraphs[0]; nos(pc); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER
    b_ra.seek(0); pc.add_run().add_picture(b_ra,width=Inches(3.4),height=Inches(2.1))

    # Inner breakdown table — columns sum to 4100 DXA (fits inside 4320 cell with margins)
    col_w_n=[1080,760,760,760,740]   # total = 4100
    inner_w=sum(col_w_n)
    months_hdr=["Jan","Feb","Mar","Total"]

    inner=doc.add_table(rows=1,cols=5)
    inner.alignment=WD_TABLE_ALIGNMENT.LEFT
    set_tbl_w(inner,inner_w)

    r0=inner.rows[0]
    hc(r0.cells[0],"",col_w_n[0])
    for j,m in enumerate(months_hdr):
        hc(r0.cells[j+1],m,col_w_n[j+1],fill=H["teal"] if m=="Total" else H["navy"])

    inner.add_row(); r1=inner.rows[1]
    dc(r1.cells[0],"PT Referrals",col_w_n[0],fill=H["lblue"],bold=True,tcol=NAVY)
    for j,v in enumerate(ref_row): dc(r1.cells[j+1],v,col_w_n[j+1],bold=(j==3),tcol=NAVY if j==3 else None)

    inner.add_row(); r2=inner.rows[2]
    dc(r2.cells[0],"Assessments",col_w_n[0],fill=H["lgreen"],bold=True,tcol=GREEN)
    for j,v in enumerate(ass_row): dc(r2.cells[j+1],v,col_w_n[j+1],bold=(j==3),tcol=GREEN if j==3 else None)

    # Nest inner table into right cell; keep one trailing paragraph (Word requires it)
    tbl_cell._tc.append(inner._tbl)
    for p_el in list(tbl_cell._tc.findall(qn("w:p"))):
        tbl_cell._tc.remove(p_el)
    tbl_cell._tc.append(OxmlElement("w:p"))   # required trailing paragraph

    gap(doc,50)

    # ── PROGRAMS + THERAPY + WORKFORCE ───────────────────────────────────────
    sec(doc,"Program Mix, Therapy Delivery and Workforce"); gap(doc,44)

    # Row 1: Program donut (left) + Therapy & Workforce (right, stacked)
    L=4500; R=4860
    t4=mtbl(doc,2,[L,R])
    lc4=t4.cell(0,0); rc4=t4.cell(0,1)
    cell_setup(lc4,L,borders=NB(),margins=(20,20,20,40))
    cell_setup(rc4,R,borders=NB(),margins=(20,20,40,20))

    # Left: donut chart
    p=lc4.paragraphs[0]; nos(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    b_prog.seek(0); p.add_run().add_picture(b_prog,width=Inches(2.9),height=Inches(2.6))

    # Right: therapy + workforce side by side inside a nested structure
    # Use a sub-table in right cell
    def add_img_to_cell(cell, buf, wi, hi):
        p=cell.paragraphs[0]; nos(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        buf.seek(0); p.add_run().add_picture(buf,width=Inches(wi),height=Inches(hi))

    # Add therapy chart
    p2=rc4.paragraphs[0]; nos(p2); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    b_min.seek(0); p2.add_run().add_picture(b_min,width=Inches(3.1),height=Inches(1.5))
    # Add small gap
    pg=rc4.add_paragraph(); sp(pg,14,0)
    # Add workforce chart
    p3=rc4.add_paragraph(); nos(p3); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    b_work.seek(0); p3.add_run().add_picture(b_work,width=Inches(3.1),height=Inches(1.5))

    gap(doc,20)

    # Caption row: 3 columns
    cw3=TW//3
    captions=[
        f"**Ambulation + Strength/Balance** {programs['ambulation']} ({round(programs['ambulation']/prog_total*100)}%), "
        f"**Chest Physio / Pain Modality** {programs['wt_bearing']}, **Strengthening + ROM** {programs['strengthening']}, "
        f"**AAROM/PROM** {programs['arom_prom']} residents.",
        f"**{therapy['one2one']:,} min** 1:1 therapy (~{hrs_1on1} hrs). "
        f"**{therapy['evaluation']:,} min** evaluations (~{hrs_eval} hrs). "
        f"**{therapy['group_sessions']} group sessions** per week.",
        f"PTA: **{staffing['pta_1on1']}h/wk** 1:1 + **{staffing['pta_group']}h/wk** group = **{staffing['pta_total']}h/wk** total. "
        f"PT: **{staffing['pt_hours']}h/wk** (Mon, Tue, Thu).",
    ]
    t5=mtbl(doc,3,[cw3,cw3,cw3])
    for ci,cap in enumerate(captions):
        c=t5.cell(0,ci)
        lb={"style":"single","sz":4,"color":H["teal"]} if ci>0 else None
        cell_setup(c,cw3,borders=dict(top=None,bottom=None,right=None,left=lb),
                   margins=(0,0,140 if ci>0 else 0,100))
        p=c.paragraphs[0]; nos(p)
        for j,seg in enumerate(cap.split("**")):
            if seg: rn(p,seg,bold=(j%2==1),sz=8.5,color=MID)

    gap(doc,45)
    p=doc.add_paragraph(); sp(p,20,0); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rn(p,f"End of Report   |   {quarter}   |   {home_name}   |   A & G Physiotherapy Inc.",
       sz=8,color=MID,italic=True)

    safe_home = re.sub(r"[^\w]+", "_", home_name).strip("_")
    out=os.path.join(os.path.dirname(os.path.abspath(__file__)),
                     f"PT_{quarter.replace(' ','_')}_{safe_home}_Report.docx")
    doc.save(out); print(f"Saved:  {out}")
    print("Exporting PDF...")
    pdf = to_pdf(out)
    if pdf: print(f"Saved:  {pdf}")
    return out

if __name__=="__main__":
    if len(sys.argv)>=2:
        excel_path=sys.argv[1]
    else:
        data_dir=os.path.join(os.path.dirname(os.path.abspath(__file__)),"..","data")
        candidates=sorted(
            [os.path.join(data_dir,f) for f in os.listdir(data_dir)
             if f.endswith((".xlsx",".xls")) and not f.startswith("~")],
            key=os.path.getmtime, reverse=True)
        if not candidates:
            print("No Excel file found in ../data/ — pass the path as an argument."); sys.exit(1)
        excel_path=candidates[0]
        print(f"Using: {os.path.basename(excel_path)}")
    generate(excel_path)
